# import textract
# text = textract.process("Quiz1-9.docx", "r" )

from docx import Document



def reverse(word):
    re_word = "".join(reversed(word))
    return re_word

doc = Document("Quiz1-9.docx")
strr=""


for para in doc.paragraphs:
    for word in range(len(para)):   
        tmp = word.decode('UTF-8')
        if tmp.isdigit():
            tmp = reverse(tmp)
        strr += tmp
    b_strr=strr.encode('UTF-8')
#print(b_strr)

f = open("demofile.docx", "wb")
f.write(b_strr)
f.close()

#print(type(numbers[1]).decode('UTF-8'))



# import docx2txt
# result = docx2txt.process("Quiz1-9.docx")

# a=result[4:14]

# #document = Document()
# mystyle = a.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)

# run.style = mystyle
# font = run.font
# font.rtl = True